"""
Convert AMASAMYA's markdown plan documents into accessible .docx files.

Why this exists: JAWS/NVDA users navigate Word documents by heading
level (H key), list type (L key for lists, T for tables, etc). That
only works if the document uses Word's built-in semantic styles
('Heading 1', 'List Bullet', etc.) — not just larger or bolded text.
A naive markdown-to-docx pipeline (e.g., pandoc with default settings)
produces visually-correct but semantically-empty documents.

This converter:
  * Maps Markdown #/##/### to Heading 1/2/3 (semantic, AT-navigable).
  * Maps - and * bullet lines to Word's 'List Bullet' style.
  * Maps "1." numbered lines to 'List Number'.
  * Maps GFM tables to real Word tables with a header row marked
    via repeatHeader (XML twiddle — python-docx doesn't expose it
    natively).
  * Sets the document title in core properties so AT announces it on
    open ("Tester Brief — AMASAMYA, by Akhilesh Malani").
  * Sets w:lang on the default style so AT picks up English without
    guessing.
  * Renders inline `code` as the Code Char style (real Char style,
    not just monospace font) so AT announces "code" semantics where
    supported.

Usage:
    python beta/md_to_docx.py            # converts all known files
    python beta/md_to_docx.py FILE.md    # converts one
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt


REPO_ROOT = Path(__file__).resolve().parent.parent

# Source -> destination + document title (what JAWS/NVDA will announce
# when the file is opened).
DOCS = [
    ("beta/tester-brief.md",
     "beta/tester-brief.docx",
     "AMASAMYA Beta — Tester Brief"),
    ("beta/feedback-form.md",
     "beta/feedback-form.docx",
     "AMASAMYA Beta — Feedback Form Question Set"),
    ("beta/outreach-emails.md",
     "beta/outreach-emails.docx",
     "AMASAMYA Beta — Outreach Email Templates"),
    ("beta/pre-flight-checklist.md",
     "beta/pre-flight-checklist.docx",
     "AMASAMYA Beta — Pre-Flight Checklist"),
    ("beta/funding/strategy.md",
     "beta/funding/strategy.docx",
     "AMASAMYA — Beta-Stage Funding Strategy"),
    ("beta/funding/one-pager.md",
     "beta/funding/one-pager.docx",
     "AMASAMYA — One-Page Brief for Funders"),
    ("beta/funding/email-templates.md",
     "beta/funding/email-templates.docx",
     "AMASAMYA — Funding Outreach Email Templates"),
    ("beta/funding/week-one-action-plan.md",
     "beta/funding/week-one-action-plan.docx",
     "AMASAMYA — Funding Week-One Action Plan"),
    ("CONTRIBUTORS.md",
     "CONTRIBUTORS.docx",
     "AMASAMYA — Contributors"),
    ("beta/README.md",
     "beta/README.docx",
     "AMASAMYA Beta — Document Index"),
]


# ───────────────────────────────────────────────────────────────────
# Inline-formatting helper
# ───────────────────────────────────────────────────────────────────

# Order matters: code first (so **inside `code`** isn't mis-bolded),
# then bold (**…**), then italic (*…*), then links ([text](url)).
INLINE_PATTERNS = [
    (re.compile(r"`([^`]+)`"), "code"),
    (re.compile(r"\*\*([^*]+)\*\*"), "bold"),
    (re.compile(r"(?<![*\w])\*([^*]+)\*(?!\w)"), "italic"),
    (re.compile(r"\[([^\]]+)\]\(([^)]+)\)"), "link"),
]


def _split_inline(text: str):
    """Yield (kind, payload) tokens for a string of inline markdown.

    Returns 'plain', 'code', 'bold', 'italic', or 'link' tokens.
    Link payload is (text, url); others are str.
    """
    if not text:
        return
    # Walk the string, peeling off the earliest match of any pattern.
    pos = 0
    while pos < len(text):
        earliest = None
        for pat, kind in INLINE_PATTERNS:
            m = pat.search(text, pos)
            if m and (earliest is None or m.start() < earliest[0].start()):
                earliest = (m, kind)
        if earliest is None:
            yield ("plain", text[pos:])
            return
        m, kind = earliest
        if m.start() > pos:
            yield ("plain", text[pos:m.start()])
        if kind == "link":
            yield ("link", (m.group(1), m.group(2)))
        else:
            yield (kind, m.group(1))
        pos = m.end()


def _apply_inline(paragraph, text: str, code_style):
    """Add runs to `paragraph` from inline-formatted markdown text."""
    for kind, payload in _split_inline(text):
        if kind == "plain":
            paragraph.add_run(payload)
        elif kind == "bold":
            run = paragraph.add_run(payload)
            run.bold = True
        elif kind == "italic":
            run = paragraph.add_run(payload)
            run.italic = True
        elif kind == "code":
            run = paragraph.add_run(payload)
            run.font.name = "Consolas"
            # Apply the Code Char style if present.
            if code_style is not None:
                run.style = code_style
        elif kind == "link":
            txt, url = payload
            _add_hyperlink(paragraph, txt, url)


def _add_hyperlink(paragraph, text, url):
    """Insert a real Word hyperlink (not just blue underlined text).

    AT announces "link" + the visible text; JAWS users can press
    Enter on it to follow.
    """
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    # Inherit from Hyperlink Char style if it exists; falls back to default colour otherwise.
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


# ───────────────────────────────────────────────────────────────────
# Block-level parser (line-by-line; keeps the converter readable)
# ───────────────────────────────────────────────────────────────────

HEADING_RE   = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
HRULE_RE     = re.compile(r"^\s*---+\s*$")
LI_BULLET_RE = re.compile(r"^(\s*)[-*]\s+(.+)$")
LI_NUM_RE    = re.compile(r"^(\s*)(\d+)\.\s+(.+)$")
TASK_RE      = re.compile(r"^\s*[-*]\s+\[([ xX])\]\s+(.+)$")
TABLE_SEP_RE = re.compile(r"^\s*\|?\s*[:\- ]+\|[:\- |]+\|?\s*$")
CODEBLOCK_RE = re.compile(r"^```")
BLOCKQUOTE_RE = re.compile(r"^>\s?(.*)$")


def _ensure_doc_styles(doc):
    """Make sure the styles we reference actually exist."""
    styles = doc.styles
    # Code Char style — used for inline `code`. Word ships 'HTML Code'
    # in some templates; create our own if missing.
    if "Inline Code" not in [s.name for s in styles]:
        try:
            cs = styles.add_style("Inline Code", WD_STYLE_TYPE.CHARACTER)
            cs.font.name = "Consolas"
            cs.font.size = Pt(10)
        except Exception:
            cs = None
    return styles["Inline Code"] if "Inline Code" in [s.name for s in styles] else None


def _set_doc_title(doc, title):
    """Set the document's core title — what JAWS/NVDA announce on open."""
    cp = doc.core_properties
    cp.title = title
    cp.language = "en-US"
    cp.author = "Akhilesh Malani"


def _set_doc_language(doc, lang="en-US"):
    """Force xml:lang on the document defaults so AT doesn't guess."""
    styles_element = doc.styles.element
    rpr_default = styles_element.find(qn("w:docDefaults"))
    if rpr_default is None:
        return
    rpr = rpr_default.find(qn("w:rPrDefault"))
    if rpr is None:
        return
    inner = rpr.find(qn("w:rPr"))
    if inner is None:
        inner = OxmlElement("w:rPr")
        rpr.append(inner)
    lang_el = inner.find(qn("w:lang"))
    if lang_el is None:
        lang_el = OxmlElement("w:lang")
        inner.append(lang_el)
    lang_el.set(qn("w:val"), lang)
    lang_el.set(qn("w:eastAsia"), lang)
    lang_el.set(qn("w:bidi"), lang)


def _add_table(doc, rows):
    """Insert a Word table from a list-of-list-of-cell-strings.

    First row is treated as the header row (bold + repeat-on-page-break).
    """
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Light Grid Accent 1"  # accessible default theme
    # Mark the first row as a header row that repeats on page breaks.
    first_row = table.rows[0]
    trPr = first_row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)
    for r_idx, row_cells in enumerate(rows):
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""  # clear default empty paragraph
            p = cell.paragraphs[0]
            text = row_cells[c_idx] if c_idx < len(row_cells) else ""
            if r_idx == 0:
                run = p.add_run(text.strip())
                run.bold = True
            else:
                _apply_inline(p, text.strip(), None)


def _flush_table(doc, table_buf):
    if not table_buf:
        return
    rows = []
    for line in table_buf:
        # Split GFM table row, strip leading/trailing pipes.
        line = line.strip()
        if line.startswith("|"):
            line = line[1:]
        if line.endswith("|"):
            line = line[:-1]
        cells = [c.strip() for c in line.split("|")]
        rows.append(cells)
    _add_table(doc, rows)


def convert(md_path: Path, docx_path: Path, title: str):
    text = md_path.read_text(encoding="utf-8")
    doc = Document()
    _set_doc_title(doc, title)
    _set_doc_language(doc)
    code_style = _ensure_doc_styles(doc)

    # First-page document title — gives sighted readers what AT
    # already gets via core properties.
    h = doc.add_paragraph(title, style="Title")

    in_code = False
    code_buffer: list[str] = []
    table_buffer: list[str] = []
    in_blockquote = False

    def flush_blockquote(quote_lines):
        if not quote_lines:
            return
        p = doc.add_paragraph(style="Intense Quote")
        _apply_inline(p, " ".join(quote_lines), code_style)

    blockquote_lines: list[str] = []

    for raw in text.splitlines():
        # Code block toggling
        if CODEBLOCK_RE.match(raw):
            if in_code:
                # End of code block — drop in as preformatted paragraph
                code_text = "\n".join(code_buffer)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = "Consolas"
                run.font.size = Pt(10)
                code_buffer = []
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_buffer.append(raw)
            continue

        # Tables: collect contiguous pipe lines, skip the separator row.
        if "|" in raw and (table_buffer or raw.strip().startswith("|") or raw.count("|") >= 2):
            if TABLE_SEP_RE.match(raw):
                continue  # separator row — skip
            if raw.strip() == "":
                _flush_table(doc, table_buffer)
                table_buffer = []
                continue
            table_buffer.append(raw)
            continue
        elif table_buffer:
            _flush_table(doc, table_buffer)
            table_buffer = []

        # Blockquote
        bq_match = BLOCKQUOTE_RE.match(raw)
        if bq_match:
            blockquote_lines.append(bq_match.group(1))
            in_blockquote = True
            continue
        elif in_blockquote and raw.strip() == "":
            flush_blockquote(blockquote_lines)
            blockquote_lines = []
            in_blockquote = False
            continue
        elif in_blockquote:
            flush_blockquote(blockquote_lines)
            blockquote_lines = []
            in_blockquote = False

        # Empty line — paragraph break, no-op
        if raw.strip() == "":
            continue

        # Horizontal rule
        if HRULE_RE.match(raw):
            p = doc.add_paragraph()
            p.add_run("―" * 30)
            continue

        # Heading
        m = HEADING_RE.match(raw)
        if m:
            level = len(m.group(1))
            heading_text = m.group(2)
            # Strip Markdown emphasis from heading text — Word headings
            # take the style from the paragraph, so embedded **bold** is
            # noisy. Just use the plain text.
            heading_text = re.sub(r"[*_`]", "", heading_text)
            doc.add_heading(heading_text, level=min(level, 9))
            continue

        # Task list (- [ ] or - [x])
        tm = TASK_RE.match(raw)
        if tm:
            checked = tm.group(1).lower() == "x"
            body = tm.group(2)
            p = doc.add_paragraph(style="List Bullet")
            box = p.add_run("☒  " if checked else "☐  ")
            box.font.name = "Segoe UI Symbol"
            _apply_inline(p, body, code_style)
            continue

        # Bullet list
        bm = LI_BULLET_RE.match(raw)
        if bm:
            indent = len(bm.group(1)) // 2
            body = bm.group(2)
            style_name = "List Bullet" if indent == 0 else "List Bullet 2"
            try:
                p = doc.add_paragraph(style=style_name)
            except KeyError:
                p = doc.add_paragraph(style="List Bullet")
            _apply_inline(p, body, code_style)
            continue

        # Numbered list
        nm = LI_NUM_RE.match(raw)
        if nm:
            body = nm.group(3)
            try:
                p = doc.add_paragraph(style="List Number")
            except KeyError:
                p = doc.add_paragraph()
                p.add_run(f"{nm.group(2)}. ")
            _apply_inline(p, body, code_style)
            continue

        # Plain paragraph
        p = doc.add_paragraph()
        _apply_inline(p, raw, code_style)

    # Flush trailing buffers
    _flush_table(doc, table_buffer)
    if blockquote_lines:
        flush_blockquote(blockquote_lines)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)
    try:
        rel_src = md_path.relative_to(REPO_ROOT)
        rel_dst = docx_path.relative_to(REPO_ROOT)
        print(f"  OK  {rel_src}  ->  {rel_dst}")
    except ValueError:
        print(f"  OK  {md_path.name}  ->  {docx_path.name}")


def main():
    if len(sys.argv) > 1:
        # Single-file mode for iterative testing.
        src = Path(sys.argv[1]).resolve()
        dst = src.with_suffix(".docx")
        title = f"AMASAMYA — {src.stem.replace('-', ' ').title()}"
        convert(src, dst, title)
        return
    print("Converting AMASAMYA plan documents to accessible .docx …")
    for src_rel, dst_rel, title in DOCS:
        src = REPO_ROOT / src_rel
        dst = REPO_ROOT / dst_rel
        if not src.exists():
            print(f"  WARN  {src_rel} not found - skipping")
            continue
        convert(src, dst, title)
    print("Done.")


if __name__ == "__main__":
    main()
